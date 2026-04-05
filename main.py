import sys
import os
import json
import base64
import time
import tempfile
import subprocess
import re as _re

# Globalny hook dla nieobsłużonych wyjątków (widoczny w .exe bez konsoli)
def _excepthook(etype, value, tb):
    import traceback
    msg = "".join(traceback.format_exception(etype, value, tb))
    try:
        from PySide6.QtWidgets import QApplication, QMessageBox
        _a = QApplication.instance() or QApplication(sys.argv)
        QMessageBox.critical(None, "Błąd krytyczny", msg)
    except Exception:
        pass
    sys.__excepthook__(etype, value, tb)
    sys.exit(1)
sys.excepthook = _excepthook

import openpyxl
from dotenv import dotenv_values

__version__ = "1.0.5"
GITHUB_REPO = "pjotermartwica/ShiftFlow"

def _collect_env_paths():
    seen_dirs = set()
    env_paths = []

    def _add_dir(path):
        if not path:
            return
        abs_path = os.path.abspath(path)
        if not os.path.isdir(abs_path) or abs_path in seen_dirs:
            return
        seen_dirs.add(abs_path)
        env_paths.append(os.path.join(abs_path, ".env"))

    if getattr(sys, 'frozen', False):
        exe_dir = os.path.dirname(os.path.abspath(sys.executable))
        _add_dir(exe_dir)
        _add_dir(os.path.dirname(exe_dir))

    _add_dir(os.getcwd())
    _add_dir(os.path.dirname(os.path.abspath(__file__)))

    if sys.argv and sys.argv[0]:
        argv_dir = os.path.dirname(os.path.abspath(sys.argv[0]))
        _add_dir(argv_dir)
        _add_dir(os.path.dirname(argv_dir))

    local_appdata = os.environ.get("LOCALAPPDATA", "").strip()
    if local_appdata:
        _add_dir(os.path.join(local_appdata, "ShiftFlow"))

    return env_paths


_ENV_SEARCH_PATHS = _collect_env_paths()
_PREFERRED_ENV_PATH = _ENV_SEARCH_PATHS[0] if _ENV_SEARCH_PATHS else os.path.join(os.getcwd(), ".env")


def _refresh_gemini_api_key():
    current_key = os.environ.get("GEMINI_API_KEY", "").strip()
    if current_key:
        return current_key, None, []

    invalid_env_paths = []
    for env_path in _ENV_SEARCH_PATHS:
        if not os.path.isfile(env_path):
            continue
        try:
            env_values = dotenv_values(env_path)
        except Exception:
            invalid_env_paths.append(env_path)
            continue

        file_key = (env_values.get("GEMINI_API_KEY") or "").strip()
        if file_key:
            os.environ["GEMINI_API_KEY"] = file_key
            return file_key, env_path, invalid_env_paths

        invalid_env_paths.append(env_path)

    return "", None, invalid_env_paths

# google.genai jest importowany leniwie w AIWorker.run() żeby uniknąć crash w PyInstaller
genai = None
genai_client = None
from openpyxl.styles import PatternFill, Font
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from PySide6.QtWidgets import (QApplication, QMainWindow, QTableView, QVBoxLayout,
                             QWidget, QPushButton, QDockWidget, QTextEdit, QLineEdit,
                             QHeaderView, QFileDialog, QMessageBox, QLabel, QTabWidget,
                             QStatusBar, QMenu, QColorDialog, QInputDialog,
                             QProgressDialog, QShortcut)
from PySide6.QtCore import Qt, QAbstractTableModel, QModelIndex, QThread, Signal, QSize, QPoint, QRect
from PySide6.QtGui import QColor, QFont, QUndoStack, QUndoCommand, QPainter, QKeyEvent, QKeySequence
from PySide6.QtNetwork import QNetworkAccessManager, QNetworkRequest, QNetworkReply
from PySide6.QtCore import QUrl
import qdarkstyle

# --- KONFIGURACJA GEMINI ---
# Klucz API jest opcjonalny przy starcie aplikacji; wymagany dopiero przy użyciu AI.
_api_key, _loaded_env_path, _invalid_env_paths = _refresh_gemini_api_key()
if _api_key:
    os.environ["GEMINI_API_KEY"] = _api_key


def _load_app_config() -> dict:
    """Wczytaj config.json z katalogu EXE lub CWD."""
    for candidate in [
        os.path.join(os.path.dirname(os.path.abspath(
            sys.executable if getattr(sys, 'frozen', False) else __file__
        )), "config.json"),
        os.path.join(os.getcwd(), "config.json"),
    ]:
        if os.path.isfile(candidate):
            try:
                with open(candidate, "r", encoding="utf-8") as _f:
                    return json.load(_f)
            except Exception:
                pass
    return {}


_APP_CONFIG: dict = _load_app_config()
_GRAFIK_FILTER = "Pliki grafiku (*.grafik);;JSON (*.json);;Wszystkie pliki (*)"

# --- KONFIGURACJA DNI I LOKALIZACJI ---
DAYS_CONFIG = {
    "Poniedziałek": ["SP-1", "SP-2", "SP-1/2"],
    "Wtorek":       ["SP-1", "SP-2", "SP-1/2"],
    "Środa":        ["SP-1", "SP-2", "SP-1/2"],
    "Czwartek":     ["SP-1", "SP-2", "SP-1/2"],
    "Piątek":       ["SP-1", "SP-2", "SP-1/2"],
}
HEADER_SEP = "|"  # Separator w nagłówkach złożonych: "Dzień|Lokalizacja"

# --- WARSTWA PAMIĘCI TRWAŁEJ (HIPOKAMP) ---
class BrainMemory:
    def __init__(self, file_path="brain_memory.json"):
        self.file_path = file_path
        self.preferences = {}
        self.hour_algorithm = "standard"
        self.load()

    def load(self):
        if os.path.exists(self.file_path):
            try:
                with open(self.file_path, "r", encoding="utf-8") as f:
                    data = json.load(f)
                    self.preferences = data.get("preferences", {})
                    self.hour_algorithm = data.get("hour_algorithm", "standard")
            except Exception as e:
                print(f"Błąd ładowania pamięci mózgowej: {e}")

    def save(self):
        data = {
            "preferences": self.preferences,
            "hour_algorithm": self.hour_algorithm
        }
        with open(self.file_path, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)

    def set_preference(self, key, value):
        self.preferences[str(key)] = value
        self.save()

    def get_preferences(self):
        return dict(self.preferences)

    def set_hour_algorithm(self, algorithm_name):
        self.hour_algorithm = algorithm_name
        self.save()

    def get_hour_algorithm(self):
        return self.hour_algorithm


# --- MODEL DANYCH: KOMÓRKA Z KOLOREM ---
class Cell:
    def __init__(self, value="", color="#FFFFFF"):
        self.value = str(value)
        self.color = color
    
    def to_dict(self):
        return {"value": self.value, "color": self.color}


# --- KOMENDA HISTORYCZNA ---
class ModelStateCommand(QUndoCommand):
    def __init__(self, model, old_state, new_state, description="Zmiana modelu"):
        super().__init__(description)
        self.model = model
        self.old_state = old_state
        self.new_state = new_state
        # QUndoStack.push() wywołuje redo() od razu; pomijamy pierwszy raz,
        # bo stan "new_state" jest już zastosowany przez akcję użytkownika.
        self._first_redo = True

    def undo(self):
        self.model.set_state(self.old_state)

    def redo(self):
        if self._first_redo:
            self._first_redo = False
            return
        self.model.set_state(self.new_state)


# --- NAGŁÓWEK HIERARCHICZNY (DZIEŃ → LOKALIZACJA) ---
class HierarchicalHeaderView(QHeaderView):
    """Dwupoziomowy nagłówek: grupuje kolumny po dniu, lokalizacje poniżej.

    Nagłówki zawierające HEADER_SEP ('|') są traktowane jako złożone: 'Dzień|Lokalizacja'.
    Nagłówki bez separatora zajmują pełną wysokość.
    """

    def __init__(self, parent=None):
        super().__init__(Qt.Horizontal, parent)
        self.setDefaultAlignment(Qt.AlignCenter)
        self.setSectionsClickable(True)
        self.setHighlightSections(True)
        self.sectionResized.connect(lambda *_: self.viewport().update())

    # --- helpers ---

    def _is_compound(self):
        m = self.model()
        if not m:
            return False
        for i in range(m.columnCount()):
            h = m.headerData(i, Qt.Horizontal, Qt.DisplayRole)
            if h and HEADER_SEP in str(h):
                return True
        return False

    def _groups(self):
        """Zwraca [(group_name | None, first_logical_col, count), ...]"""
        m = self.model()
        if not m:
            return []
        result = []
        prev = object()
        for col in range(m.columnCount()):
            h = str(m.headerData(col, Qt.Horizontal, Qt.DisplayRole) or "")
            g = h.split(HEADER_SEP, 1)[0].strip() if HEADER_SEP in h else None
            if g == prev and result:
                result[-1] = (result[-1][0], result[-1][1], result[-1][2] + 1)
            else:
                result.append((g, col, 1))
            prev = g
        return result

    # --- reimplementacja ---

    def sizeHint(self):
        base = super().sizeHint()
        if self._is_compound():
            base.setHeight(max(base.height(), 24) * 2)
        return base

    def paintEvent(self, event):
        if not self._is_compound():
            super().paintEvent(event)
            return

        painter = QPainter(self.viewport())
        m = self.model()
        if not m:
            painter.end()
            return

        h = self.height()
        th = h // 2
        bh = h - th

        bg_group = QColor("#2d5f8a")
        bg_sub = QColor("#1a3a5c")
        bg_flat = QColor("#2a2a2a")
        fg = QColor("#FFFFFF")
        border = QColor("#444444")

        groups = self._groups()

        # --- dolny rząd: lokalizacje lub proste nagłówki ---
        for col in range(m.columnCount()):
            if self.isSectionHidden(col):
                continue
            x = self.sectionViewportPosition(col)
            w = self.sectionSize(col)
            header = str(m.headerData(col, Qt.Horizontal, Qt.DisplayRole) or "")

            if HEADER_SEP in header:
                sub = header.split(HEADER_SEP, 1)[1].strip()
                r = QRect(x, th, w, bh)
                painter.fillRect(r, bg_sub)
                painter.setPen(fg)
                painter.drawText(r, Qt.AlignCenter, sub)
                painter.setPen(border)
                painter.drawLine(r.topLeft(), r.topRight())
                painter.drawLine(r.right(), r.top(), r.right(), r.bottom())
                painter.drawLine(r.bottomLeft(), r.bottomRight())
            else:
                r = QRect(x, 0, w, h)
                painter.fillRect(r, bg_flat)
                f = painter.font()
                f.setBold(True)
                painter.setFont(f)
                painter.setPen(fg)
                painter.drawText(r, Qt.AlignCenter, header)
                f.setBold(False)
                painter.setFont(f)
                painter.setPen(border)
                painter.drawRect(r.adjusted(0, 0, -1, -1))

        # --- górny rząd: grupy (dni) ---
        for gname, start, cnt in groups:
            if gname is None:
                continue
            x = self.sectionViewportPosition(start)
            tw = sum(
                self.sectionSize(start + i)
                for i in range(cnt)
                if not self.isSectionHidden(start + i)
            )
            if tw <= 0:
                continue
            r = QRect(x, 0, tw, th)
            painter.fillRect(r, bg_group)
            f = painter.font()
            f.setBold(True)
            painter.setFont(f)
            painter.setPen(fg)
            painter.drawText(r, Qt.AlignCenter, gname)
            f.setBold(False)
            painter.setFont(f)
            painter.setPen(border)
            painter.drawRect(r.adjusted(0, 0, -1, -1))

        painter.end()


# --- MODEL TABELI ---
class ScheduleTableModel(QAbstractTableModel):
    state_about_to_change = Signal(dict)
    state_changed = Signal(dict)

    def __init__(self):
        super().__init__()
        self.headers = ["Pracownik"]
        self.rows = []
        self.spans = []  # [[r, c, row_span, col_span], ...]
        self.load_from_file()

    def get_state(self):
        return {
            "headers": list(self.headers),
            "rows": [[cell.to_dict() for cell in row] for row in self.rows],
            "spans": [list(s) for s in self.spans]
        }

    def set_state(self, state):
        saved_headers = state.get("headers", ["Pracownik"])
        saved_rows = state.get("rows", [])

        self.beginResetModel()
        self.headers = list(saved_headers)
        self.rows = []
        for row_data in saved_rows:
            row = [Cell(cell.get("value", ""), cell.get("color", "#FFFFFF"))
                   for cell in row_data]
            self.rows.append(row)
        self.spans = [list(s) for s in state.get("spans", [])]
        self.endResetModel()

        self.save_to_file()
        self.state_changed.emit(self.get_state())

    def load_from_file(self):
        """Wczytaj dane z harmonogram.json"""
        if os.path.exists("harmonogram.json"):
            try:
                with open("harmonogram.json", "r", encoding="utf-8") as f:
                    data = json.load(f)
                    self.headers = data.get("headers", ["Pracownik"])
                    self.rows = []
                    for row_data in data.get("rows", []):
                        row = [Cell(cell.get("value", ""), cell.get("color", "#FFFFFF"))
                               for cell in row_data]
                        self.rows.append(row)
                    self.spans = [list(s) for s in data.get("spans", [])]
            except Exception as e:
                print(f"Błąd wczytywania: {e}")

    def save_to_file(self):
        """Zapisz dane do harmonogram.json"""
        data = self.get_state()
        with open("harmonogram.json", "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)

    def rowCount(self, parent=QModelIndex()):
        return len(self.rows)

    def columnCount(self, parent=QModelIndex()):
        return len(self.headers)

    def data(self, index, role=Qt.DisplayRole):
        if not index.isValid():
            return None

        row, col = index.row(), index.column()
        if row >= len(self.rows) or col >= len(self.rows[row]):
            return None

        cell = self.rows[row][col]

        if role == Qt.DisplayRole or role == Qt.EditRole:
            return cell.value
        elif role == Qt.BackgroundRole:
            return QColor(cell.color) if QColor(cell.color).isValid() else QColor("#FFFFFF")
        elif role == Qt.ForegroundRole:
            # Auto-contrast: biały tekst na ciemnym tle, czarny na jasnym
            bg = QColor(cell.color) if QColor(cell.color).isValid() else QColor("#FFFFFF")
            luminance = 0.299 * bg.red() + 0.587 * bg.green() + 0.114 * bg.blue()
            return QColor("#000000") if luminance > 140 else QColor("#FFFFFF")
        elif role == Qt.FontRole:
            if col == 0:
                font = QFont()
                font.setBold(True)
                return font

        return None

    def setData(self, index, value, role=Qt.EditRole):
        if not index.isValid():
            return False

        row, col = index.row(), index.column()
        if row >= len(self.rows) or col >= len(self.headers):
            return False

        if role == Qt.EditRole:
            old_state = self.get_state()
            self.state_about_to_change.emit(old_state)

            self.rows[row][col].value = str(value)
            self.dataChanged.emit(index, index)
            self.save_to_file()
            self.state_changed.emit(self.get_state())
            return True

        return False

    def headerData(self, section, orientation, role=Qt.DisplayRole):
        if role == Qt.DisplayRole:
            if orientation == Qt.Horizontal:
                return self.headers[section] if section < len(self.headers) else ""
            else:
                return str(section + 1)
        return None

    def flags(self, index):
        return Qt.ItemIsSelectable | Qt.ItemIsEnabled | Qt.ItemIsEditable

    def add_worker(self):
        """Dodaj nowy wiersz - pracownika"""
        self.state_about_to_change.emit(self.get_state())
        num_workers = len(self.rows) + 1
        new_row = [Cell(f"Pracownik {num_workers}") if i == 0 
                   else Cell("") for i in range(len(self.headers))]
        self.beginInsertRows(QModelIndex(), len(self.rows), len(self.rows))
        self.rows.append(new_row)
        self.endInsertRows()
        self.save_to_file()
        self.state_changed.emit(self.get_state())

    def remove_worker(self, row):
        """Usuń wiersz i przepisz numery"""
        if 0 <= row < len(self.rows):
            self.state_about_to_change.emit(self.get_state())
            self.beginRemoveRows(QModelIndex(), row, row)
            self.rows.pop(row)
            self.endRemoveRows()
            self._reindex_workers()
            self.save_to_file()
            self.state_changed.emit(self.get_state())

    def _reindex_workers(self):
        """Przepisz numery pracowników TYLKO jeśli komórka ma domyślną nazwę 'Pracownik N'"""
        import re
        for i, row in enumerate(self.rows):
            if row and re.fullmatch(r'Pracownik \d+', row[0].value):
                row[0].value = f"Pracownik {i + 1}"

    def add_day(self):
        """Dodaj nową kolumnę - dzień"""
        self.state_about_to_change.emit(self.get_state())
        day_num = len(self.headers)
        self.headers.append(f"Dzień {day_num}")

        for row in self.rows:
            row.append(Cell(""))

        self.layoutChanged.emit()
        self.save_to_file()
        self.state_changed.emit(self.get_state())

    def remove_day(self, col):
        """Usuń kolumnę - dzień"""
        if col > 0 and col < len(self.headers):
            self.state_about_to_change.emit(self.get_state())
            self.headers.pop(col)
            for row in self.rows:
                if col < len(row):
                    row.pop(col)
            self.layoutChanged.emit()
            self.save_to_file()
            self.state_changed.emit(self.get_state())

    def rename_header(self, col, name):
        """Zmień nazwę nagłówka kolumny"""
        if 0 <= col < len(self.headers):
            self.state_about_to_change.emit(self.get_state())
            self.headers[col] = str(name)
            self.headerDataChanged.emit(Qt.Horizontal, col, col)
            self.save_to_file()
            self.state_changed.emit(self.get_state())

    # ------------------------------------------------------------------
    # NOWE NARZĘDZIA – bezpośredni dostęp po współrzędnych
    # ------------------------------------------------------------------

    def update_cell(self, r, c, val, color=None):
        """Atomic update of cell (r, c). Used by AI tool-calling."""
        if not (0 <= r < len(self.rows) and 0 <= c < len(self.rows[r])):
            return False
        self.state_about_to_change.emit(self.get_state())
        self.rows[r][c].value = str(val)
        if color:
            self.rows[r][c].color = color
        idx = self.index(r, c)
        self.dataChanged.emit(idx, idx)
        self.save_to_file()
        self.state_changed.emit(self.get_state())
        return True

    def delete_range(self, r1, c1, r2, c2):
        """Clear all cells in the rectangle [r1..r2, c1..c2]."""
        self.state_about_to_change.emit(self.get_state())
        for r in range(max(0, r1), min(len(self.rows), r2 + 1)):
            for c in range(max(0, c1), min(len(self.rows[r]), c2 + 1)):
                self.rows[r][c].value = ""
                self.rows[r][c].color = "#FFFFFF"
        self.layoutChanged.emit()
        self.save_to_file()
        self.state_changed.emit(self.get_state())

    def insert_rows(self, at, count=1):
        """Insert 'count' blank rows before position 'at'."""
        at = max(0, min(at, len(self.rows)))
        self.state_about_to_change.emit(self.get_state())
        self.beginInsertRows(QModelIndex(), at, at + count - 1)
        for i in range(count):
            new_row = [Cell("") for _ in range(len(self.headers))]
            new_row[0].value = f"Pracownik {at + i + 1}"
            self.rows.insert(at + i, new_row)
        self.endInsertRows()
        self._reindex_workers()
        self.save_to_file()
        self.state_changed.emit(self.get_state())

    def insert_columns(self, at, count=1, names=None):
        """Insert 'count' blank columns before position 'at'."""
        at = max(1, min(at, len(self.headers)))
        self.state_about_to_change.emit(self.get_state())
        for i in range(count):
            col_idx = at + i
            col_name = (names[i] if names and i < len(names)
                        else f"Dzień {col_idx}")
            self.headers.insert(col_idx, col_name)
            for row in self.rows:
                row.insert(col_idx, Cell(""))
        self.layoutChanged.emit()
        self.save_to_file()
        self.state_changed.emit(self.get_state())

    def remove_rows_range(self, r1, r2):
        """Remove rows r1..r2 inclusive."""
        r1 = max(0, r1)
        r2 = min(len(self.rows) - 1, r2)
        if r1 > r2:
            return
        self.state_about_to_change.emit(self.get_state())
        self.beginRemoveRows(QModelIndex(), r1, r2)
        del self.rows[r1:r2 + 1]
        self.endRemoveRows()
        self._reindex_workers()
        self.save_to_file()
        self.state_changed.emit(self.get_state())

    def remove_columns_range(self, c1, c2):
        """Remove columns c1..c2 inclusive."""
        c1 = max(0, c1)
        c2 = min(len(self.headers) - 1, c2)
        if c1 > c2:
            return
        self.state_about_to_change.emit(self.get_state())
        del self.headers[c1:c2 + 1]
        for row in self.rows:
            del row[c1:c2 + 1]
        # Invalidate spans that overlap removed columns
        self.spans = [
            s for s in self.spans
            if not (s[1] >= c1 and s[1] <= c2)
        ]
        self.layoutChanged.emit()
        self.save_to_file()
        self.state_changed.emit(self.get_state())

    def set_span(self, r, c, row_span, col_span):
        """Register a merge span and persist it."""
        # Remove any existing span starting at same cell
        self.spans = [s for s in self.spans if not (s[0] == r and s[1] == c)]
        if row_span > 1 or col_span > 1:
            self.spans.append([r, c, row_span, col_span])
        self.save_to_file()

    # ------------------------------------------------------------------
    # KOLUMNY ZŁOŻONE: DZIEŃ → LOKALIZACJA
    # ------------------------------------------------------------------

    def generate_schedule_columns(self, days_config):
        """Wygeneruj kolumny z konfiguracji {Dzień: [Lokalizacje]}.

        Kolumna 0 to 'Pracownik', reszta to 'Dzień|Lokalizacja'.
        """
        self.state_about_to_change.emit(self.get_state())
        self.beginResetModel()

        new_headers = ["Pracownik"]
        for day, locations in days_config.items():
            for loc in locations:
                new_headers.append(f"{day}{HEADER_SEP}{loc}")

        self.headers = new_headers
        for row in self.rows:
            while len(row) < len(self.headers):
                row.append(Cell(""))
            while len(row) > len(self.headers):
                row.pop()
        self.spans = []

        self.endResetModel()
        self.save_to_file()
        self.state_changed.emit(self.get_state())

    def resolve_column(self, day, location):
        """Zamień (dzień, lokalizacja) na fizyczny indeks kolumny. -1 jeśli nie znaleziono."""
        target = f"{day}{HEADER_SEP}{location}"
        for i, h in enumerate(self.headers):
            if h == target:
                return i
        # fuzzy: case-insensitive
        target_l = target.lower()
        for i, h in enumerate(self.headers):
            if h.lower() == target_l:
                return i
        return -1

    def get_column_map(self):
        """Zwraca słownik (dzień, lokalizacja) → indeks kolumny."""
        mapping = {}
        for i, h in enumerate(self.headers):
            if HEADER_SEP in h:
                parts = h.split(HEADER_SEP, 1)
                mapping[(parts[0].strip(), parts[1].strip())] = i
        return mapping

    def apply_commands(self, commands):
        """Zastosuj listę komend od AI — jako jeden blok undo z jednym zapisem."""
        if not commands:
            return
        self.state_about_to_change.emit(self.get_state())

        # Blokujemy sygnały stanowe w sub-metodach, żeby nie tworzyły
        # osobnych wpisów undo (cała lista = jeden undo).
        saved = self.blockSignals(True)
        try:
            self._apply_commands_inner(commands)
        finally:
            self.blockSignals(saved)

        self.beginResetModel()
        self.endResetModel()
        self.save_to_file()
        self.state_changed.emit(self.get_state())

    def _apply_commands_inner(self, commands):
        """Wewnętrzna pętla komend — bez sygnałów."""
        for cmd in commands:
            action = cmd.get("action")

            if action == "add_worker":
                num = len(self.rows) + 1
                new_row = [Cell(f"Pracownik {num}") if i == 0 
                           else Cell("") for i in range(len(self.headers))]
                self.rows.append(new_row)
            elif action == "remove_worker":
                row = cmd.get("row", 0)
                if 0 <= row < len(self.rows):
                    self.rows.pop(row)
                    self._reindex_workers()
            elif action == "add_day":
                day_num = len(self.headers)
                self.headers.append(f"Dzień {day_num}")
                for row in self.rows:
                    row.append(Cell(""))
            elif action == "remove_day":
                col = cmd.get("col", 1)
                if 0 < col < len(self.headers):
                    self.headers.pop(col)
                    for row in self.rows:
                        if col < len(row):
                            row.pop(col)
            elif action == "edit":
                row = cmd.get("row")
                col = cmd.get("col")
                day = cmd.get("day")
                location = cmd.get("location")
                if day and location:
                    col = self.resolve_column(day, location)
                    if col < 0:
                        continue
                value = cmd.get("value", "")
                color = cmd.get("color", "#FFFFFF")
                if row is not None and col is not None:
                    # auto-extend if AI references rows/cols beyond current size
                    while row >= len(self.rows):
                        self.rows.append([Cell("") for _ in range(len(self.headers))])
                    while col >= len(self.headers):
                        self.headers.append(f"Kol {len(self.headers)}")
                        for r in self.rows:
                            r.append(Cell(""))
                    self.rows[row][col].value = value
                    self.rows[row][col].color = color
            elif action == "update_cell":
                r = cmd.get("r", cmd.get("row", 0))
                day = cmd.get("day")
                location = cmd.get("location")
                if day and location:
                    c = self.resolve_column(day, location)
                    if c < 0:
                        continue
                else:
                    c = cmd.get("c", cmd.get("col", 0))
                val = cmd.get("val", cmd.get("value", ""))
                color = cmd.get("color")
                while r >= len(self.rows):
                    self.rows.append([Cell("") for _ in range(len(self.headers))])
                while c >= len(self.headers):
                    self.headers.append(f"Kol {len(self.headers)}")
                    for rr in self.rows:
                        rr.append(Cell(""))
                self.rows[r][c].value = str(val)
                if color:
                    self.rows[r][c].color = color
            elif action == "delete_range":
                r1, c1 = cmd.get("r1", 0), cmd.get("c1", 0)
                r2, c2 = cmd.get("r2", 0), cmd.get("c2", 0)
                for r in range(max(0, r1), min(len(self.rows), r2 + 1)):
                    for c in range(max(0, c1), min(len(self.rows[r]), c2 + 1)):
                        self.rows[r][c].value = ""
                        self.rows[r][c].color = "#FFFFFF"
            elif action == "insert_row":
                at = max(0, min(cmd.get("at", len(self.rows)), len(self.rows)))
                count = cmd.get("count", 1)
                for i in range(count):
                    new_row = [Cell("") for _ in range(len(self.headers))]
                    new_row[0].value = f"Pracownik {at + i + 1}"
                    self.rows.insert(at + i, new_row)
                self._reindex_workers()
            elif action == "insert_col":
                at = max(1, min(cmd.get("at", len(self.headers)), len(self.headers)))
                count = cmd.get("count", 1)
                names = cmd.get("names")
                for i in range(count):
                    col_idx = at + i
                    col_name = (names[i] if names and i < len(names)
                                else f"Dzień {col_idx}")
                    self.headers.insert(col_idx, col_name)
                    for row in self.rows:
                        row.insert(col_idx, Cell(""))
            elif action == "remove_rows":
                r1 = max(0, cmd.get("r1", 0))
                r2 = min(len(self.rows) - 1, cmd.get("r2", 0))
                if r1 <= r2:
                    del self.rows[r1:r2 + 1]
                    self._reindex_workers()
            elif action == "remove_cols":
                c1 = max(1, cmd.get("c1", 1))
                c2 = min(len(self.headers) - 1, cmd.get("c2", 1))
                if c1 <= c2:
                    del self.headers[c1:c2 + 1]
                    for row in self.rows:
                        del row[c1:c2 + 1]
                    self.spans = [s for s in self.spans
                                  if not (s[1] >= c1 and s[1] <= c2)]
            elif action == "merge_cells":
                self.set_span(
                    cmd.get("r", 0), cmd.get("c", 0),
                    cmd.get("row_span", 1), cmd.get("col_span", 1)
                )
            elif action == "rename_header":
                col = cmd.get("col", 0)
                if 0 <= col < len(self.headers):
                    self.headers[col] = str(cmd.get("name", ""))
            elif action == "generate_schedule":
                config = cmd.get("config", {})
                if config:
                    new_headers = ["Pracownik"]
                    for day_name, locations in config.items():
                        for loc in locations:
                            new_headers.append(f"{day_name}{HEADER_SEP}{loc}")
                    self.headers = new_headers
                    for row in self.rows:
                        while len(row) < len(self.headers):
                            row.append(Cell(""))
                        while len(row) > len(self.headers):
                            row.pop()
                    self.spans = []

# --- WĄTEK AI: AGENTA AUTONOMICZNY ---
class AIWorker(QThread):
    response_ready = Signal(list)
    error_occurred = Signal(str)
    
    def __init__(self, prompt, current_data, imported_data=None, imported_image=None):
        super().__init__()
        self.prompt = prompt
        self.current_data = current_data
        self.imported_data = imported_data
        self.imported_image = imported_image
    
    def run(self):
        try:
            _api_key, _loaded_env_path, _invalid_env_paths = _refresh_gemini_api_key()
            if not _api_key:
                if _invalid_env_paths:
                    checked_path = _invalid_env_paths[0]
                    message = (
                        "Brak poprawnego klucza GEMINI_API_KEY.\n\n"
                        f"Znaleziono plik .env w: {checked_path}\n"
                        "ale nie zawiera on poprawnej wartości GEMINI_API_KEY.\n\n"
                        f"Popraw plik i wpisz: GEMINI_API_KEY=twój_klucz\n"
                        "Następnie spróbuj ponownie."
                    )
                else:
                    checked_locations = "\n".join(f"- {path}" for path in _ENV_SEARCH_PATHS[:5])
                    message = (
                        "Brak klucza GEMINI_API_KEY.\n\n"
                        "Aplikacja szukała pliku .env w lokalizacjach:\n"
                        f"{checked_locations}\n\n"
                        f"Utwórz plik .env w: {_PREFERRED_ENV_PATH}\n"
                        "z treścią: GEMINI_API_KEY=twój_klucz\n\n"
                        "Następnie spróbuj ponownie."
                    )
                self.error_occurred.emit(
                    message
                )
                return

            # Lazy import google.genai (unika crash w PyInstaller)
            try:
                import google.genai as genai
            except ModuleNotFoundError:
                try:
                    import google.generativeai as genai
                except ModuleNotFoundError:
                    self.error_occurred.emit("Brak pakietu google.genai; pip install google-genai")
                    return

            # Configure SDK
            try:
                genai.configure(api_key=_api_key)
            except AttributeError:
                pass  # new SDK doesn't use configure()

            # Detect which SDK is being used
            _use_new_sdk = not hasattr(genai, 'GenerativeModel')

            system_instruction = """
Jesteś Autonomicznym Menedżerem Grafiku Pracy. Operujesz WYŁĄCZNIE przez narzędzia JSON.
Zwracaj TYLKO czysty JSON jako listę komend, bez żadnego tekstu poza JSON.

=== DOSTĘPNE NARZĘDZIA (action) ===

Podstawowe:
  {"action": "edit",        "row": R, "col": C, "value": "TEKST", "color": "#HEX"}
  {"action": "update_cell", "r": R,   "c": C,   "val":   "TEKST", "color": "#HEX"}
  {"action": "update_cell", "day": "Poniedziałek", "location": "SP-1", "row": R, "val": "TEKST", "color": "#HEX"}

Struktura grafiku (hierarchiczna):
  {"action": "generate_schedule", "config": {"Poniedziałek": ["SP-1","SP-2","SP-1/2"], "Wtorek": ["SP-1","SP-2","SP-1/2"], ...}}

Struktura wierszy/kolumn:
  {"action": "add_worker"}
  {"action": "remove_worker", "row": R}
  {"action": "add_day"}
  {"action": "remove_day", "col": C}
  {"action": "insert_row",  "at": R, "count": N}
  {"action": "insert_col",  "at": C, "count": N, "names": ["Dzień 1", ...]}
  {"action": "remove_rows", "r1": R1, "r2": R2}
  {"action": "remove_cols", "c1": C1, "c2": C2}

Zakres i scalanie:
  {"action": "delete_range", "r1": R1, "c1": C1, "r2": R2, "c2": C2}
  {"action": "merge_cells",  "r": R, "c": C, "row_span": RS, "col_span": CS}
  {"action": "rename_header", "col": C, "name": "Nowa nazwa"}

Preferencje i UI:
  {"action": "set_preference", "key": "nazwa", "value": "wartość"}
  {"action": "universal_execute", "command": "add_tab", "name": "Nowa zakładka"}
  {"action": "universal_execute", "command": "set_hour_algorithm", "algorithm": "pauzy"}
  {"action": "universal_execute", "command": "set_ui_theme", "theme": "light"}

=== LOGIKA ZASTĘPSTW (SUBSTITUTION) ===
Dane overtime w current_data["overtime"]:
  { "Nazwisko": { "row_index": R, "worked_h": H, "ratio_pct": P } }

Algorytm przypisania zastępstwa:
1. Oblicz ratio = (worked_h * 100) / standard_h dla każdego pracownika.
2. Wybierz pracownika z NAJNIŻSZYM ratio_pct jako zastępcę.
3. Jeśli prompt zezwala na nadgodziny (ratio > 100%), przypisz mimo to i dodaj sekcję
   {"action": "set_preference", "key": "overtime_log", "value": "NAZWISKO +Xh"}
4. Nigdy nie przypisuj tej samej osoby dwa razy w tej samej kolumnie (dniu).

=== OCR ZE ZDJĘCIA ===
Jeśli dołączono zdjęcie:
- Odczytaj WSZYSTKIE osoby, miejsca, dni i godziny (OCR).
- Kolumna 0 = imię i nazwisko, kolejne = dni tygodnia.
- Wartość komórki = zakres godzin ("8:00-16:00") lub puste jeśli wolne.
- Usuń powtórzenia (ta sama osoba max raz na dzień/slot).
- Posortuj kolumny chronologicznie.
- Kolory: stosuj paletę poniżej.

=== PALETA KOLORÓW ===
Kategorie zmian i statusów:
  Wczesna zmiana (rano):    #FFD700 (złoty) lub #FFF176 (jasny żółty)
  Zmiana dzienna:           #64B5F6 (błękitny) lub #81D4FA (jasnoniebieski)
  Zmiana popołudniowa:      #FF8C00 (pomarańczowy) lub #FFB74D (łososiowy)
  Zmiana nocna:             #9370DB (fioletowy) lub #CE93D8 (lawendowy)
  Wolne / urlop:            #A5D6A7 (jasnozielony) lub #C8E6C9 (miętowy)
  Świetlica / dyżur:       #4DB6AC (morski) lub #80CBC4 (turkusowy)
  Zastępstwo:               #EF9A9A (różowy) lub #F48FB1 (malinowy)
  Nagłówki / etykiety:      #B0BEC5 (szary stalowy) lub #CFD8DC (jasnoszary)
  Nadgodziny (ratio>100%):  #E57373 (czerwony pastelowy)

Zasady kolorowania:
- Każda kategoria zmian powinna mieć JEDNORODNY kolor w obrębie grafiku.
- Stosuj kolory z palety powyżej — nie wymyślaj własnych.
- Kolumna 0 (imiona) opcjonalnie koloruj na #B0BEC5 dla odróżnienia od danych.
- Puste komórki zostawiaj bez koloru (domyślny motyw).
- Jeśli grafik ma nagłówek/tytuł, użyj merge_cells + kolor #B0BEC5.
- Stosuj alternujące odcienie dla sąsiednich zmian tego samego typu (np. #FFD700 i #FFF176).

=== ZASADY OGÓLNE ===
- Kolumna 0 = imiona/nazwiska pracowników (row_index odpowiada pozycji w rows[]).
- Współrzędne: row = indeks wiersza (0-based), col = indeks kolumny (0-based).
- Możesz też adresować po (day, location): {"action":"update_cell","day":"Poniedziałek","location":"SP-1","row":0,"val":"Anna"}.
- Optymalny grafik: równomierne rozłożenie, brak konfliktów, ratio < 100%.
- Twórz estetyczne, czytelne grafiki — dobierz kolory, scalaj nagłówki, wyrównaj dane.
- Użyj generate_schedule aby stworzyć strukturę Dzień→Lokalizacja przed wypełnieniem danych.

Przykład (hierarchiczny grafik):
[
  {"action": "generate_schedule", "config": {"Poniedziałek": ["SP-1","SP-2","SP-1/2"], "Wtorek": ["SP-1","SP-2","SP-1/2"]}},
  {"action": "insert_row", "at": 0, "count": 3},
  {"action": "update_cell", "day": "Poniedziałek", "location": "SP-1", "row": 0, "val": "8:00-14:00", "color": "#FFD700"},
  {"action": "update_cell", "day": "Poniedziałek", "location": "SP-2", "row": 0, "val": "9:00-15:00", "color": "#FFF176"},
  {"action": "update_cell", "r": 0, "c": 0, "val": "Anna Kowalska", "color": "#B0BEC5"},
  {"action": "update_cell", "r": 1, "c": 0, "val": "Jan Nowak", "color": "#B0BEC5"},
  {"action": "update_cell", "day": "Wtorek", "location": "SP-1", "row": 1, "val": "14:00-22:00", "color": "#FF8C00"}
]
"""
            
            imported_section = ""
            if self.imported_data:
                lines = ["\t".join(row) for row in self.imported_data]
                imported_section = (
                    "\n\nDane z zaimportowanego pliku (tabela z prawdziwymi nazwiskami pracowników):\n"
                    + "\n".join(lines)
                )

            image_instruction = ""
            if self.imported_image:
                image_instruction = (
                    "\n\nDo zapytania dołączono ZDJĘCIE tabeli grafiku. "
                    "Odczytaj z niego WSZYSTKIE osoby, miejsca pracy, dni i godziny. "
                    "Wygeneruj komendy aby odtworzyć tę tabelę w aplikacji. "
                    "Pamiętaj o usunięciu powtórzeń i posortowaniu kolumn chronologicznie."
                )

            full_prompt = f"""{system_instruction}

Obecny stan grafiku:
{json.dumps(self.current_data, ensure_ascii=False, indent=2)}{imported_section}{image_instruction}

Życzenie użytkownika:
{self.prompt}

Zwróć TYLKO JSON jako listę komend, bez żadnych wyjaśnień."""

            _IMAGE_MIME = {
                '.jpg': 'image/jpeg', '.jpeg': 'image/jpeg',
                '.png': 'image/png', '.gif': 'image/gif',
                '.webp': 'image/webp', '.bmp': 'image/bmp',
            }

            response = None
            last_exc = None
            for _attempt in range(4):   # up to 4 tries: 0s, 5s, 15s, 30s wait
                if _attempt > 0:
                    wait_sec = [5, 15, 30][_attempt - 1]
                    time.sleep(wait_sec)
                try:
                    if self.imported_image:
                        ext = os.path.splitext(self.imported_image)[1].lower()
                        mime = _IMAGE_MIME.get(ext, 'image/jpeg')
                        with open(self.imported_image, 'rb') as _f:
                            img_bytes = _f.read()

                        if _use_new_sdk:
                            from google.genai import Client, types
                            _client = Client(api_key=os.environ.get("GEMINI_API_KEY"))
                            response = _client.models.generate_content(
                                model='gemini-2.5-flash',
                                contents=[
                                    types.Part.from_bytes(data=img_bytes, mime_type=mime),
                                    full_prompt
                                ]
                            )
                        else:
                            image_part = {
                                "inline_data": {
                                    "mime_type": mime,
                                    "data": base64.b64encode(img_bytes).decode()
                                }
                            }
                            model = genai.GenerativeModel('gemini-2.5-flash')
                            response = model.generate_content([image_part, full_prompt])
                    elif _use_new_sdk:
                        from google.genai import Client
                        _client = Client(api_key=os.environ.get("GEMINI_API_KEY"))
                        response = _client.models.generate_content(
                            model='gemini-2.5-flash',
                            contents=full_prompt
                        )
                    else:
                        model = genai.GenerativeModel('gemini-2.5-flash')
                        response = model.generate_content(full_prompt)
                    break  # success
                except Exception as _e:
                    last_exc = _e
                    _err_str = str(_e)
                    if "503" not in _err_str and "UNAVAILABLE" not in _err_str and "quota" not in _err_str.lower():
                        raise  # non-retryable, fail immediately
                    # else retry
            if response is None:
                raise last_exc
            text_res = response.text.strip()
            
            # Spróbuj wyodrębnić JSON
            if "```json" in text_res:
                text_res = text_res.split("```json")[1].split("```")[0].strip()
            elif "```" in text_res:
                text_res = text_res.split("```")[1].split("```")[0].strip()
            
            commands = json.loads(text_res)
            if not isinstance(commands, list):
                commands = [commands]
            
            self.response_ready.emit(commands)
        
        except json.JSONDecodeError as e:
            self.error_occurred.emit(f"Błąd parsowania JSON: {e}")
        except Exception as e:
            err = str(e)
            # 503 / UNAVAILABLE — retry with backoff
            if "503" in err or "UNAVAILABLE" in err or "quota" in err.lower():
                self.error_occurred.emit(
                    f"Błąd AI: serwer przeciążony (503). Spróbuj ponownie za chwilę.\nSzczegóły: {err}"
                )
            else:
                self.error_occurred.emit(f"Błąd AI: {err}")

# --- AUTO-UPDATER ---
class ScriptUpdater(QThread):
    """
    Pobiera nową wersję main.py z GitHub Raw i zapisuje ją do
    %LOCALAPPDATA%/ShiftFlow/. Bootstrap załaduje ją przy kolejnym starcie.
    Nie wymaga przebudowy .exe.
    """
    # (new_version)  — nowa wersja została pobrana
    updated   = Signal(str)
    # (current_version) — jesteśmy aktualni
    up_to_date = Signal(str)
    failed    = Signal(str)

    _GITHUB_RAW = f"https://raw.githubusercontent.com/{GITHUB_REPO}/main"
    _APP_DIR = os.path.join(
        os.environ.get("LOCALAPPDATA", os.path.expanduser("~")), "ShiftFlow"
    )

    def run(self):
        import urllib.request
        raw = self._GITHUB_RAW
        try:
            # 1. Sprawdź wersję
            req = urllib.request.Request(
                f"{raw}/version.txt",
                headers={"User-Agent": "ShiftFlow-Updater"},
            )
            with urllib.request.urlopen(req, timeout=8) as r:
                remote_ver = r.read().decode().strip()

            def _parts(v: str) -> list:
                return [int(x) for x in _re.findall(r"\d+", v)]

            if not (_parts(remote_ver) > _parts(__version__)):
                self.up_to_date.emit(__version__)
                return

            # 2. Pobierz main.py (~100 KB)
            req2 = urllib.request.Request(
                f"{raw}/main.py",
                headers={"User-Agent": "ShiftFlow-Updater"},
            )
            with urllib.request.urlopen(req2, timeout=30) as r:
                data = r.read()

            # 3. Zapisz atomowo
            os.makedirs(self._APP_DIR, exist_ok=True)
            dest = os.path.join(self._APP_DIR, "main.py")
            tmp  = dest + ".tmp"
            with open(tmp, "wb") as f:
                f.write(data)
            os.replace(tmp, dest)
            with open(os.path.join(self._APP_DIR, "version.txt"), "w") as f:
                f.write(remote_ver)

            self.updated.emit(remote_ver)
        except Exception as e:
            self.failed.emit(str(e))


# --- GŁÓWNE OKNO APLIKACJI ---
class ScheduleApp(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("ShiftFlow - AI Work Scheduler")
        self._apply_startup_geometry()
        self.setStyleSheet(qdarkstyle.load_stylesheet())

        self._project_file: str | None = None
        self._unsaved = False

        self.model = ScheduleTableModel()
        self.ai_worker = None
        self.brain_memory = BrainMemory()
        self._imported_file_data = None
        self._imported_image_path = None

        self.undo_stack = QUndoStack(self)
        self.undo_stack.setUndoLimit(200)
        self._history_enabled = True
        self._pending_state = None

        self.model.state_about_to_change.connect(self.on_model_about_to_change)
        self.model.state_changed.connect(self.on_model_state_changed)
        self.model.state_changed.connect(self._mark_dirty)

        self._setup_ui()

        self.statusBar().showMessage("System gotowy")
        self.set_ai_status("AI: gotowe", "#00FF00")
        # Bootstrap obsługuje auto-aktualizację main.py przy każdym uruchomieniu.
        # Tutaj przechowujemy referencje do manualne sprawdzenia w panelu.
        self._script_updater: ScriptUpdater | None = None
    
    def _setup_ui(self):
        """Zbuduj interfejs użytkownika"""
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        layout = QVBoxLayout(central_widget)

        # Wskaźnik Pasek stanu AI
        self.live_status_label = QLabel("Analiza pamięci...")
        self.live_status_label.setStyleSheet("font-weight: bold; color: #00FF00;")
        layout.addWidget(self.live_status_label)

        # Zakładki (wielopłaszczyznowy widok)
        self.tab_widget = QTabWidget()
        self.tab_widget.setTabsClosable(True)
        self.tab_widget.tabCloseRequested.connect(self.on_tab_close_requested)

        self.table_view = QTableView()
        self._h_header = HierarchicalHeaderView(self.table_view)
        self.table_view.setHorizontalHeader(self._h_header)
        self.table_view.setModel(self.model)
        self.table_view.setEditTriggers(QTableView.DoubleClicked | QTableView.SelectedClicked | QTableView.EditKeyPressed | QTableView.AnyKeyPressed)
        self._h_header.setSectionResizeMode(QHeaderView.Stretch)
        self.table_view.verticalHeader().setSectionResizeMode(QHeaderView.ResizeToContents)
        self.table_view.setSelectionMode(QTableView.ExtendedSelection)
        self.table_view.setSelectionBehavior(QTableView.SelectItems)
        self.table_view.setContextMenuPolicy(Qt.CustomContextMenu)
        self.table_view.customContextMenuRequested.connect(self.on_table_context_menu)
        # Rename column header by double-clicking on it
        self._h_header.sectionDoubleClicked.connect(self._rename_header_dialog)
        self.table_view.installEventFilter(self)

        table_container = QWidget()
        table_layout = QVBoxLayout(table_container)
        table_layout.addWidget(self.table_view)

        self.tab_widget.addTab(table_container, "Tabela główna")
        layout.addWidget(self.tab_widget)
        
        # Boczny panel (DockWidget)
        dock = QDockWidget("Panel Kontrolny", self)
        dock_widget = QWidget()
        dock_layout = QVBoxLayout(dock_widget)

        # ── PLIK / ZAPIS I ODCZYT ─────────────────────────────────────────
        dock_layout.addWidget(QLabel("--- PLIK ---"))

        btn_new = QPushButton("📄  Nowy grafik")
        btn_new.clicked.connect(self.new_project)
        dock_layout.addWidget(btn_new)

        btn_save = QPushButton("💾  Zapisz  (Ctrl+S)")
        btn_save.clicked.connect(self.save_project)
        dock_layout.addWidget(btn_save)

        btn_save_as = QPushButton("💾  Zapisz jako…  (Ctrl+Shift+S)")
        btn_save_as.clicked.connect(self.save_project_as)
        dock_layout.addWidget(btn_save_as)

        btn_open = QPushButton("📂  Otwórz…  (Ctrl+O)")
        btn_open.clicked.connect(self.open_project)
        dock_layout.addWidget(btn_open)

        # Skróty klawiaturowe
        QShortcut(QKeySequence("Ctrl+S"), self).activated.connect(self.save_project)
        QShortcut(QKeySequence("Ctrl+Shift+S"), self).activated.connect(self.save_project_as)
        QShortcut(QKeySequence("Ctrl+O"), self).activated.connect(self.open_project)
        QShortcut(QKeySequence("Ctrl+N"), self).activated.connect(self.new_project)

        # ── SIATA / WIERSZE I KOLUMNY ────────────────────────────────────
        dock_layout.addWidget(QLabel("--- SIATKA ---"))

        btn_add_row = QPushButton("➕ Dodaj wiersz")
        btn_add_row.clicked.connect(lambda: self.perform_action(self.model.add_worker))
        dock_layout.addWidget(btn_add_row)

        btn_add_col = QPushButton("➕ Dodaj kolumnę…")
        btn_add_col.clicked.connect(self._add_column_dialog)
        btn_add_col.setToolTip("Wpisz nazwę nowej kolumny")
        dock_layout.addWidget(btn_add_col)

        btn_del_sel = QPushButton("✂  Usuń zaznaczone wiersze")
        btn_del_sel.clicked.connect(self._delete_selected_rows)
        dock_layout.addWidget(btn_del_sel)

        btn_del_col = QPushButton("✂  Usuń zaznaczone kolumny")
        btn_del_col.clicked.connect(self._delete_selected_cols)
        dock_layout.addWidget(btn_del_col)

        btn_rename_h = QPushButton("✏  Zmień nazwę kolumny…")
        btn_rename_h.clicked.connect(lambda: self._rename_header_dialog(
            self.table_view.currentIndex().column()))
        dock_layout.addWidget(btn_rename_h)

        # Szybkie zestawy kolumn
        dock_layout.addWidget(QLabel("Szybkie zestawy:"))
        btn_preset_hours = QPushButton("🕐  Godziny lekcyjne (1–9)")
        btn_preset_hours.clicked.connect(lambda: self._insert_preset(
            ["1", "2", "3", "4", "5", "6", "7", "8", "9"]))
        dock_layout.addWidget(btn_preset_hours)

        btn_preset_sp = QPushButton("🏫  SP1 + SP2 (świetlice)")
        btn_preset_sp.clicked.connect(lambda: self._insert_preset(["SP-1", "SP-2"]))
        dock_layout.addWidget(btn_preset_sp)

        btn_preset_days = QPushButton("📅  Pon–Pt (5 dni)")
        btn_preset_days.clicked.connect(lambda: self._insert_preset(
            ["Poniedziałek", "Wtorek", "Środa", "Czwartek", "Piątek"]))
        dock_layout.addWidget(btn_preset_days)

        btn_gen_schedule = QPushButton("📋  Generuj strukturę grafiku")
        btn_gen_schedule.clicked.connect(self._generate_schedule_dialog)
        btn_gen_schedule.setToolTip("Dzień → Lokalizacje (SP-1, SP-2, SP-1/2) × 5 dni")
        dock_layout.addWidget(btn_gen_schedule)

        btn_add_tab = QPushButton("📑  Dodaj zakładkę")
        btn_add_tab.clicked.connect(lambda: self.perform_action(self.add_tab))
        dock_layout.addWidget(btn_add_tab)

        btn_undo = QPushButton("↶  Wstecz (Undo)")
        btn_undo.clicked.connect(self.undo)
        dock_layout.addWidget(btn_undo)
        self.btn_undo = btn_undo

        btn_redo = QPushButton("↷  Do przodu (Redo)")
        btn_redo.clicked.connect(self.redo)
        dock_layout.addWidget(btn_redo)
        self.btn_redo = btn_redo

        # Sekcja AI
        dock_layout.addWidget(QLabel("\n--- AGENT AI ---"))
        dock_layout.addWidget(QLabel("Wpisz zadanie:"))
        
        self.prompt_input = QTextEdit()
        self.prompt_input.setPlaceholderText("Np. 'Stwórz grafik dla 5 osób na tydzień z kolorami'")
        self.prompt_input.setMaximumHeight(80)
        dock_layout.addWidget(self.prompt_input)

        btn_attach_file = QPushButton("📎 Załącz plik (Excel / Word / Zdjęcie)")
        btn_attach_file.clicked.connect(self.import_file_for_ai)
        btn_attach_file.setToolTip("Wczytaj Excel, Word lub zdjęcie planu jako kontekst AI")
        dock_layout.addWidget(btn_attach_file)

        self._imported_file_label = QLabel("Brak załączonego pliku")
        self._imported_file_label.setStyleSheet("color: #888888; font-style: italic;")
        dock_layout.addWidget(self._imported_file_label)
        
        btn_ask_ai = QPushButton("Zapytaj AI 🤖")
        btn_ask_ai.clicked.connect(self.ask_ai)
        btn_ask_ai.setStyleSheet("background-color: #2196F3; color: white; font-weight: bold;")
        dock_layout.addWidget(btn_ask_ai)

        self.ai_status_label = QLabel("AI: gotowe")
        self.ai_status_label.setStyleSheet("font-weight: bold; color: #00FF00;")
        dock_layout.addWidget(self.ai_status_label)
        
        btn_color_cell = QPushButton("Zmień kolor komórki 🎨")
        btn_color_cell.clicked.connect(self.pick_color_for_selected)
        dock_layout.addWidget(btn_color_cell)

        btn_sum_hours = QPushButton("Oblicz sumę godzin ⏱")
        btn_sum_hours.clicked.connect(self.add_hours_summary)
        dock_layout.addWidget(btn_sum_hours)

        # Sekcja eksportu
        dock_layout.addWidget(QLabel("\n--- EKSPORT ---"))
        
        btn_export_excel = QPushButton("Eksportuj do Excel")
        btn_export_excel.clicked.connect(self.export_to_excel)
        dock_layout.addWidget(btn_export_excel)
        
        btn_export_word = QPushButton("Eksportuj do Word")
        btn_export_word.clicked.connect(self.export_to_word)
        dock_layout.addWidget(btn_export_word)

        # Log
        dock_layout.addWidget(QLabel("\n--- STATUS ---"))
        self.status_log = QTextEdit()
        self.status_log.setReadOnly(True)
        self.status_log.setMaximumHeight(100)
        dock_layout.addWidget(self.status_log)

        btn_check_update = QPushButton("🔄  Sprawdź aktualizacje")
        btn_check_update.clicked.connect(self.check_for_updates_manual)
        dock_layout.addWidget(btn_check_update)

        dock_layout.addStretch()
        
        dock.setWidget(dock_widget)
        self.addDockWidget(Qt.RightDockWidgetArea, dock)
        
        self._log("Aplikacja uruchomiona ✓")
        self.update_undo_redo_buttons()

    def on_table_context_menu(self, pos: QPoint):
        """Rozszerzone menu PPM: kolor, scalanie, wstawianie, usuwanie, czyszczenie"""
        index = self.table_view.indexAt(pos)
        selected = self.table_view.selectedIndexes()
        menu = QMenu(self)

        # --- Kolor ---
        sec_color = menu.addSection("🎨 Kolor")
        if len(selected) > 1:
            menu.addAction(f"Zmień kolor zaznaczonych ({len(selected)})", lambda: self._apply_color_to_indexes(selected))
        if index.isValid():
            menu.addAction("Zmień kolor tej komórki", lambda: self._apply_color_to_indexes([index]))
            menu.addAction(f"Zmień kolor wiersza {index.row() + 1}", lambda: self._apply_color_to_indexes([
                self.model.index(index.row(), c) for c in range(self.model.columnCount())
            ]))

        # --- Scal / rozdziel ---
        if len(selected) > 1:
            rows_sel = sorted({i.row() for i in selected})
            cols_sel = sorted({i.column() for i in selected})
            r1, r2, c1, c2 = rows_sel[0], rows_sel[-1], cols_sel[0], cols_sel[-1]
            row_span = r2 - r1 + 1
            col_span = c2 - c1 + 1
            sec_merge = menu.addSection("⊞ Scalanie")
            menu.addAction(f"Scal zaznaczone ({row_span}×{col_span})", lambda: self._merge_selected(r1, c1, row_span, col_span))
        if index.isValid():
            menu.addAction("Rozdziel tę komórkę", lambda: self._unmerge_cell(index.row(), index.column()))

        # --- Wstawianie ---
        if index.isValid():
            sec_ins = menu.addSection("➕ Wstawianie")
            menu.addAction(f"Wstaw wiersz powyżej ({index.row() + 1})", lambda: self.perform_action(self.model.insert_rows, index.row()))
            menu.addAction(f"Wstaw wiersz poniżej ({index.row() + 2})", lambda: self.perform_action(self.model.insert_rows, index.row() + 1))
            if index.column() > 0:
                menu.addAction(f"Wstaw kolumnę z lewej ({index.column()})", lambda: self.perform_action(self.model.insert_columns, index.column()))
                menu.addAction(f"Wstaw kolumnę z prawej ({index.column() + 2})", lambda: self.perform_action(self.model.insert_columns, index.column() + 1))

        # --- Usuwanie / czyszczenie ---
        if selected:
            rows_all = sorted({i.row() for i in selected})
            cols_all = sorted({i.column() for i in selected})
            sec_del = menu.addSection("🗑 Usuwanie")
            if rows_all:
                menu.addAction(f"Usuń zaznaczone wiersze ({rows_all[0]+1}–{rows_all[-1]+1})",
                    lambda r1=rows_all[0], r2=rows_all[-1]: self.perform_action(self.model.remove_rows_range, r1, r2))
            if cols_all:
                menu.addAction(f"Usuń zaznaczone kolumny ({cols_all[0]+1}–{cols_all[-1]+1})",
                    lambda c1=cols_all[0], c2=cols_all[-1]: self.perform_action(self.model.remove_columns_range, c1, c2))
            all_rows = sorted({i.row() for i in selected})
            all_cols = sorted({i.column() for i in selected})
            menu.addAction(f"Wyczyść zakres ({len(selected)} komórek)",
                lambda r1=all_rows[0], c1=all_cols[0], r2=all_rows[-1], c2=all_cols[-1]:
                    self.perform_action(self.model.delete_range, r1, c1, r2, c2))

        menu.exec(self.table_view.viewport().mapToGlobal(pos))

    def _merge_selected(self, r, c, row_span, col_span):
        """Scal komórki i zastosuj span w widoku"""
        if r >= len(self.model.rows) or c >= len(self.model.rows[r]):
            return
        old_state = self.model.get_state()
        # Skopiuj wartość pierwszej komórki do scalenia
        val = self.model.rows[r][c].value
        color = self.model.rows[r][c].color
        # Wyczyść komórki objęte scalaniem
        for rr in range(r, r + row_span):
            for cc in range(c, c + col_span):
                if rr == r and cc == c:
                    continue
                if rr < len(self.model.rows) and cc < len(self.model.rows[rr]):
                    self.model.rows[rr][cc].value = ""
        self.model.set_span(r, c, row_span, col_span)
        self.table_view.setSpan(r, c, row_span, col_span)
        self.model.save_to_file()
        self._log(f"Scalono {row_span}×{col_span} od ({r},{c})")

    def _unmerge_cell(self, r, c):
        """Rozdziel span w (r, c)"""
        self.model.set_span(r, c, 1, 1)
        self.table_view.setSpan(r, c, 1, 1)
        self._log(f"Rozdzielono komórkę ({r},{c})")

    def _reapply_spans(self):
        """Przywróć wszystkie spany po undo/redo/load"""
        for s in self.model.spans:
            if len(s) == 4:
                self.table_view.setSpan(s[0], s[1], s[2], s[3])

    def pick_color_for_selected(self):
        """Zmień kolor zaznaczonych komórek przez przycisk w panelu"""
        selected = self.table_view.selectedIndexes()
        if not selected:
            QMessageBox.information(self, "Info",
                "Zaznacz komórki (Ctrl+klik, Shift+klik lub kliknij nagłówek wiersza).")
            return
        self._apply_color_to_indexes(selected)

    def _apply_color_to_indexes(self, indexes):
        """Otwiera QColorDialog i stosuje kolor do listy indeksów komórek"""
        if not indexes:
            return
        # Użyj koloru pierwszej zaznaczonej komórki jako punktu startowego
        first = indexes[0]
        if first.row() >= len(self.model.rows) or first.column() >= len(self.model.rows[first.row()]):
            return
        start_color = QColor(self.model.rows[first.row()][first.column()].color)
        color = QColorDialog.getColor(start_color, self, "Wybierz kolor komórek")
        if not color.isValid():
            return

        old_state = self.model.get_state()
        for idx in indexes:
            r, c = idx.row(), idx.column()
            if 0 <= r < len(self.model.rows) and 0 <= c < len(self.model.rows[r]):
                self.model.rows[r][c].color = color.name()

        # Odśwież widok
        top_left = self.model.index(
            min(i.row() for i in indexes), min(i.column() for i in indexes)
        )
        bottom_right = self.model.index(
            max(i.row() for i in indexes), max(i.column() for i in indexes)
        )
        self.model.dataChanged.emit(top_left, bottom_right)
        self.model.save_to_file()

        new_state = self.model.get_state()
        self._history_enabled = False
        try:
            self.undo_stack.push(ModelStateCommand(
                self.model, old_state, new_state, f"Zmiana koloru ({len(indexes)} komórek)"
            ))
        finally:
            self._history_enabled = True
        self.update_undo_redo_buttons()
        self._log(f"Kolor {len(indexes)} komórek zmieniony na {color.name()}")

    def add_hours_summary(self):
        """Dodaje/aktualizuje ostatnią kolumnę 'Suma godzin' – zlicza puste vs. wypełnione pola"""
        SUMMARY_HEADER = "Suma godzin"
        # Jeśli nie istnieje, dodaj kolumnę
        if SUMMARY_HEADER not in self.model.headers:
            self._pending_state = self.model.get_state()
            self.model.headers.append(SUMMARY_HEADER)
            for row in self.model.rows:
                row.append(Cell("", "#DDDDDD"))
            self.model.layoutChanged.emit()
        # Oblicz i wpisz sumy
        sum_col = self.model.headers.index(SUMMARY_HEADER)
        alg = self.brain_memory.get_hour_algorithm()
        for row_data in self.model.rows:
            filled = sum(
                1 for cell in row_data[1:sum_col] if cell.value.strip() != ""
            )
            if alg == "pauzy":
                # 8h na zmianę minus 0.5h pauza
                total = filled * 7.5
            else:
                total = filled * 8
            row_data[sum_col].value = f"{total:.1f}h"
            row_data[sum_col].color = "#B0E0E6"
        self.model.save_to_file()
        self.model.layoutChanged.emit()
        self._log(f"✓ Zaktualizowano kolumnę '{SUMMARY_HEADER}' (algorytm: {alg})")

    def on_tab_close_requested(self, index):
        if self.tab_widget.count() > 1:
            self.tab_widget.removeTab(index)
            self._log(f"Zakładka {index} zamknięta")

    def add_tab(self, name=None):
        if name is None:
            name = f"Zakładka {self.tab_widget.count() + 1}"

        table_view = QTableView()
        tab_h_header = HierarchicalHeaderView(table_view)
        table_view.setHorizontalHeader(tab_h_header)
        table_view.setModel(self.model)
        table_view.setEditTriggers(QTableView.DoubleClicked | QTableView.SelectedClicked | QTableView.EditKeyPressed | QTableView.AnyKeyPressed)
        tab_h_header.setSectionResizeMode(QHeaderView.Stretch)
        table_view.verticalHeader().setSectionResizeMode(QHeaderView.ResizeToContents)

        container = QWidget()
        inner_layout = QVBoxLayout(container)
        inner_layout.addWidget(table_view)

        self.tab_widget.addTab(container, name)
        self.tab_widget.setCurrentWidget(container)
        self._log(f"Dodano nową zakładkę: {name}")

    def execute_universal_command(self, command_data):
        command = command_data.get("command")
        if command == "add_tab":
            self.add_tab(command_data.get("name"))
            self._log("Wykonano universal_execute: add_tab")
        elif command == "set_hour_algorithm":
            alg = command_data.get("algorithm", "standard")
            self.brain_memory.set_hour_algorithm(alg)
            self._log(f"Ustawiono algorytm godzin: {alg}")
        elif command == "set_ui_theme":
            theme = command_data.get("theme", "dark")
            if theme == "light":
                self.setStyleSheet("")
            else:
                self.setStyleSheet(qdarkstyle.load_stylesheet())
            self._log(f"Zastosowano motyw UI: {theme}")
        else:
            self._log(f"Nieznana komenda universal_execute: {command}")

    def apply_ai_commands(self, commands):
        for cmd in commands:
            action = cmd.get("action")
            if action == "set_preference":
                key = cmd.get("key")
                value = cmd.get("value")
                if key is not None:
                    self.brain_memory.set_preference(key, value)
                    self._log(f"Ustawiono preferencję: {key} = {value}")
            elif action == "universal_execute":
                self.execute_universal_command(cmd)
            else:
                # standardowe działania modelu
                self.model.apply_commands([cmd])

    def on_model_about_to_change(self, state):
        if not self._history_enabled:
            return
        self._pending_state = json.loads(json.dumps(state))

    def on_model_state_changed(self, state):
        if not self._history_enabled:
            self.update_undo_redo_buttons()
            self._reapply_spans()
            return
        if self._pending_state is not None:
            old_state = self._pending_state
            self._pending_state = None
            self._history_enabled = False
            try:
                self.undo_stack.push(
                    ModelStateCommand(
                        self.model,
                        old_state,
                        json.loads(json.dumps(state)),
                        "Operacja tabeli"
                    )
                )
            finally:
                self._history_enabled = True
        self.update_undo_redo_buttons()
        self._reapply_spans()

    def update_undo_redo_buttons(self):
        if hasattr(self, 'btn_undo'):
            self.btn_undo.setEnabled(self.undo_stack.canUndo())
        if hasattr(self, 'btn_redo'):
            self.btn_redo.setEnabled(self.undo_stack.canRedo())

    def undo(self):
        if self.undo_stack.canUndo():
            self._history_enabled = False
            try:
                self.undo_stack.undo()
            finally:
                self._history_enabled = True
            self._log("↶ Cofnięto akcję (Undo)")
            self.update_undo_redo_buttons()

    def redo(self):
        if self.undo_stack.canRedo():
            self._history_enabled = False
            try:
                self.undo_stack.redo()
            finally:
                self._history_enabled = True
            self._log("↷ Przywrócono akcję (Redo)")
            self.update_undo_redo_buttons()

    def calculate_overtime(self, row):
        """Returns (worked_h, ratio_pct) for a given row index.
        ratio = (worked_h * 100) / standard_h
        Standard = 5 shifts/week * shift duration.
        """
        if not (0 <= row < len(self.model.rows)):
            return 0.0, 0.0
        summary_col = (self.model.headers.index("Suma godzin")
                       if "Suma godzin" in self.model.headers else None)
        alg = self.brain_memory.get_hour_algorithm()
        shift_h = 7.5 if alg == "pauzy" else 8.0
        standard_h = 5 * shift_h  # 5 shifts per week
        worked_shifts = sum(
            1 for c, cell in enumerate(self.model.rows[row])
            if c > 0 and c != summary_col and cell.value.strip()
        )
        worked_h = worked_shifts * shift_h
        ratio = (worked_h * 100.0 / standard_h) if standard_h else 0.0
        return round(worked_h, 1), round(ratio, 1)

    def set_ai_status(self, text, color="#00FF00"):
        self.ai_status_label.setText(text)
        self.ai_status_label.setStyleSheet(f"font-weight: bold; color: {color};")

    # ── STARTUP GEOMETRY ──────────────────────────────────────────────────

    def _apply_startup_geometry(self):
        """Ustaw rozmiar okna w oparciu o dostępne miejsce na ekranie."""
        screen = QApplication.primaryScreen()
        if screen:
            geom = screen.availableGeometry()
            cfg = _APP_CONFIG.get("window", {})
            w = int(geom.width()  * cfg.get("width_ratio",  0.85))
            h = int(geom.height() * cfg.get("height_ratio", 0.85))
            x = geom.x() + (geom.width()  - w) // 2
            y = geom.y() + (geom.height() - h) // 2
            self.setGeometry(x, y, w, h)
        else:
            self.setGeometry(100, 100, 1200, 700)

    # ── ZARZĄDZANIE PLIKIEM PROJEKTU ──────────────────────────────────────

    def _update_title(self):
        name = os.path.basename(self._project_file) if self._project_file else "Bez nazwy"
        dirty = " *" if self._unsaved else ""
        self.setWindowTitle(f"ShiftFlow — {name}{dirty}")

    def _mark_dirty(self, _state=None):
        if not self._unsaved:
            self._unsaved = True
            self._update_title()

    def new_project(self):
        if self._unsaved:
            reply = QMessageBox.question(
                self, "Niezapisane zmiany",
                "Masz niezapisane zmiany. Zapisać przed stworzeniem nowego grafiku?",
                QMessageBox.Save | QMessageBox.Discard | QMessageBox.Cancel,
            )
            if reply == QMessageBox.Save:
                self.save_project()
                if self._unsaved:   # zapis anulowany
                    return
            elif reply == QMessageBox.Cancel:
                return
        empty: dict = {"headers": ["Pracownik"], "rows": [], "spans": []}
        self._history_enabled = False
        try:
            self.model.set_state(empty)
        finally:
            self._history_enabled = True
        self.undo_stack.clear()
        self._project_file = None
        self._unsaved = False
        self._update_title()
        self._log("✓ Nowy grafik")

    def save_project(self):
        if not self._project_file:
            self.save_project_as()
            return
        self._write_project(self._project_file)

    def save_project_as(self):
        path, _ = QFileDialog.getSaveFileName(
            self, "Zapisz projekt", "", _GRAFIK_FILTER
        )
        if path:
            self._write_project(path)

    def _write_project(self, path: str):
        try:
            data = self.model.get_state()
            with open(path, "w", encoding="utf-8") as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
            self._project_file = path
            self._unsaved = False
            self._update_title()
            self._log(f"✓ Zapisano: {os.path.basename(path)}")
        except OSError as e:
            QMessageBox.critical(self, "Błąd zapisu", str(e))

    def open_project(self):
        if self._unsaved:
            reply = QMessageBox.question(
                self, "Niezapisane zmiany",
                "Masz niezapisane zmiany. Zapisać przed otwarciem innego pliku?",
                QMessageBox.Save | QMessageBox.Discard | QMessageBox.Cancel,
            )
            if reply == QMessageBox.Save:
                self.save_project()
                if self._unsaved:
                    return
            elif reply == QMessageBox.Cancel:
                return
        path, _ = QFileDialog.getOpenFileName(
            self, "Otwórz projekt", "", _GRAFIK_FILTER
        )
        if not path:
            return
        try:
            with open(path, "r", encoding="utf-8") as f:
                data = json.load(f)
            self._history_enabled = False
            try:
                self.model.set_state(data)
            finally:
                self._history_enabled = True
            self.undo_stack.clear()
            self._project_file = path
            self._unsaved = False
            self._update_title()
            self._log(f"✓ Otwarto: {os.path.basename(path)}")
        except (OSError, json.JSONDecodeError) as e:
            QMessageBox.critical(self, "Błąd otwarcia", str(e))

    def closeEvent(self, event):
        if self._unsaved:
            reply = QMessageBox.question(
                self, "Niezapisane zmiany",
                "Masz niezapisane zmiany. Zapisać przed zamknięciem?",
                QMessageBox.Save | QMessageBox.Discard | QMessageBox.Cancel,
            )
            if reply == QMessageBox.Save:
                self.save_project()
                event.accept()
            elif reply == QMessageBox.Discard:
                event.accept()
            else:
                event.ignore()
        else:
            event.accept()

    # ── AKCJE OGÓLNE ──────────────────────────────────────────────────────

    def perform_action(self, func, *args, **kwargs):
        func(*args, **kwargs)
        self.update_undo_redo_buttons()

    def ask_ai(self):
        """Poproś AI o optymalizację grafiku"""
        prompt = self.prompt_input.toPlainText().strip()
        
        if not prompt:
            QMessageBox.warning(self, "Błąd", "Proszę wpisać zadanie dla AI!")
            return
        
        self._log("⏳ Czekam na odpowiedź AI...")
        self.set_ai_status("Planowanie optymalizacji...", "#FFFF00")

        # Czytanie pamięci trwałej
        self.live_status_label.setText("Analiza pamięci...")
        memory_state = self.brain_memory.get_preferences()
        hour_algorithm = self.brain_memory.get_hour_algorithm()
        self.live_status_label.setText("Planowanie optymalizacji...")

        # Przygotuj dane do wysłania (z nadgodziną dla każdego pracownika)
        overtime_data = {}
        for i, row in enumerate(self.model.rows):
            name = row[0].value if row else f"Pracownik {i+1}"
            worked_h, ratio = self.calculate_overtime(i)
            overtime_data[name] = {"row_index": i, "worked_h": worked_h, "ratio_pct": ratio}

        current_data = {
            "headers": self.model.headers,
            "rows": [[cell.to_dict() for cell in row] for row in self.model.rows],
            "brain_memory": memory_state,
            "hour_algorithm": hour_algorithm,
            "overtime": overtime_data
        }

        self.ai_worker = AIWorker(prompt, current_data, self._imported_file_data, self._imported_image_path)
        self.ai_worker.response_ready.connect(self.on_ai_response)
        self.ai_worker.error_occurred.connect(self.on_ai_error)
        self.ai_worker.start()
    
    def on_ai_response(self, commands):
        """Obsłuż odpowiedź AI"""
        try:
            self.set_ai_status("Wdrażanie zmian...", "#00FF00")
            self.perform_action(lambda: self.apply_ai_commands(commands))
            self._log(f"✓ AI zastosowało {len(commands)} komend!")
            self.set_ai_status("AI: gotowe", "#00FF00")
        except Exception as e:
            self._log(f"❌ Błąd przy stosowaniu komend: {e}")
            self.set_ai_status("AI: błąd", "#FF0000")
    
    def on_ai_error(self, error_msg):
        """Obsłuż błąd AI"""
        self._log(f"❌ {error_msg}")
        self.set_ai_status("AI: błąd", "#FF0000")
        QMessageBox.critical(self, "Błąd", error_msg)
    
    def export_to_excel(self):
        """Eksportuj grafik do Excela"""
        file_path, _ = QFileDialog.getSaveFileName(self, "Zapisz Excel", "", "Excel Files (*.xlsx)")
        if not file_path:
            return
        
        try:
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "Grafik"
            
            # Nagłówki
            for col, header in enumerate(self.model.headers, 1):
                cell = ws.cell(row=1, column=col, value=header)
                cell.font = Font(bold=True, color="FFFFFF")
                cell.fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
            
            # Dane
            for row_idx, row in enumerate(self.model.rows, 2):
                for col_idx, cell_data in enumerate(row, 1):
                    cell = ws.cell(row=row_idx, column=col_idx, value=cell_data.value)
                    
                    # Zastosuj kolor
                    try:
                        color_hex = cell_data.color.lstrip("#")
                        cell.fill = PatternFill(start_color=color_hex, end_color=color_hex, fill_type="solid")
                    except:
                        pass
            
            # Dostosuj szerokość kolumn
            for col in ws.columns:
                ws.column_dimensions[col[0].column_letter].width = 20
            
            wb.save(file_path)
            self._log(f"✓ Grafik eksportowany do: {file_path}")
            QMessageBox.information(self, "Sukces", f"Plik zapisany: {file_path}")
        except Exception as e:
            self._log(f"❌ Błąd eksportu Excel: {e}")
            QMessageBox.critical(self, "Błąd", f"Błąd: {e}")
    
    def export_to_word(self):
        """Eksportuj grafik do Worda"""
        file_path, _ = QFileDialog.getSaveFileName(self, "Zapisz Word", "", "Word Files (*.docx)")
        if not file_path:
            return
        
        try:
            doc = Document()
            doc.add_heading("Grafik Pracy", 0)
            
            # Tabela w Wordzie
            table = doc.add_table(rows=len(self.model.rows) + 1, cols=len(self.model.headers))
            table.style = "Light Grid Accent 1"
            
            # Nagłówki
            header_cells = table.rows[0].cells
            for i, header in enumerate(self.model.headers):
                header_cells[i].text = header
            
            # Dane
            for row_idx, row in enumerate(self.model.rows, 1):
                row_cells = table.rows[row_idx].cells
                for col_idx, cell_data in enumerate(row):
                    row_cells[col_idx].text = cell_data.value
            
            doc.save(file_path)
            self._log(f"✓ Grafik eksportowany do: {file_path}")
            QMessageBox.information(self, "Sukces", f"Plik zapisany: {file_path}")
        except Exception as e:
            self._log(f"❌ Błąd eksportu Word: {e}")
            QMessageBox.critical(self, "Błąd", f"Błąd: {e}")

    # ── POMOCNICZE METODY SIATKI ──────────────────────────────────────────

    def _add_column_dialog(self):
        """Dialog: wpisz nazwę nowej kolumny i wstaw ją na końcu"""
        name, ok = QInputDialog.getText(self, "Nowa kolumna", "Nazwa kolumny:")
        if ok and name.strip():
            at = len(self.model.headers)
            self.perform_action(self.model.insert_columns, at, 1, [name.strip()])
            self._log(f"✓ Dodano kolumnę: {name.strip()}")

    def _rename_header_dialog(self, col):
        """Dialog: zmień nazwę istniejącej kolumny"""
        if col < 0 or col >= len(self.model.headers):
            return
        current = self.model.headers[col]
        name, ok = QInputDialog.getText(
            self, "Zmień nazwę kolumny", "Nowa nazwa:", text=current)
        if ok and name.strip():
            self.perform_action(self.model.rename_header, col, name.strip())
            self._log(f"✓ Kolumna {col} → \"{name.strip()}\"")

    def _delete_selected_rows(self):
        """Usuń wiersze zaznaczone w tabeli"""
        sel = self.table_view.selectionModel()
        rows = sorted({i.row() for i in sel.selectedIndexes()}, reverse=True)
        if not rows:
            QMessageBox.information(self, "Info", "Zaznacz wiersze do usunięcia.")
            return
        self.perform_action(self.model.remove_rows_range, rows[-1], rows[0])
        self._log(f"✓ Usunięto wiersze {rows[-1]+1}–{rows[0]+1}")

    def _delete_selected_cols(self):
        """Usuń kolumny zaznaczone w tabeli"""
        sel = self.table_view.selectionModel()
        cols = sorted({i.column() for i in sel.selectedIndexes()}, reverse=True)
        if not cols:
            QMessageBox.information(self, "Info",
                "Zaznacz kolumny do usunięcia.")
            return
        self.perform_action(self.model.remove_columns_range, cols[-1], cols[0])
        self._log(f"✓ Usunięto kolumny {cols[-1]+1}–{cols[0]+1}")

    def _insert_preset(self, names):
        """Wstaw zestaw gotowych kolumn na końcu tabeli"""
        at = len(self.model.headers)
        self.perform_action(self.model.insert_columns, at, len(names), names)
        self._log(f"✓ Dodano zestaw: {', '.join(names)}")

    def _generate_schedule_dialog(self):
        """Wygeneruj strukturę grafiku z DAYS_CONFIG"""
        days_info = ", ".join(f"{d} ({len(l)})" for d, l in DAYS_CONFIG.items())
        reply = QMessageBox.question(
            self, "Generuj strukturę grafiku",
            f"Wygenerować kolumny Dzień → Lokalizacja?\n\n{days_info}\n\n"
            "Obecne kolumny zostaną zastąpione.",
            QMessageBox.Yes | QMessageBox.No
        )
        if reply == QMessageBox.Yes:
            self.perform_action(self.model.generate_schedule_columns, DAYS_CONFIG)
            self._log("✓ Wygenerowano strukturę: Dzień → Lokalizacja")

    def import_file_for_ai(self):
        """Wczytaj plik Excel, Word lub zdjęcie i przekaż jako kontekst AI"""
        file_path, _ = QFileDialog.getOpenFileName(
            self, "Wybierz plik", "",
            "Obsługiwane pliki (*.xlsx *.xls *.docx *.jpg *.jpeg *.png *.gif *.webp *.bmp);;"
            "Excel (*.xlsx *.xls);;Word (*.docx);;Zdjęcia (*.jpg *.jpeg *.png *.gif *.webp *.bmp)"
        )
        if not file_path:
            return
        try:
            ext = os.path.splitext(file_path)[1].lower()
            fname = os.path.basename(file_path)
            if ext in ('.xlsx', '.xls'):
                data = self._parse_excel(file_path)
                self._imported_file_data = data
                self._imported_image_path = None
                self._imported_file_label.setText(f"📎 {fname} ({len(data)} wierszy)")
                self._imported_file_label.setStyleSheet("color: #00CC66; font-style: italic;")
                self._log(f"✓ Załadowano Excel do AI: {fname} ({len(data)} wierszy)")
            elif ext == '.docx':
                data = self._parse_word(file_path)
                self._imported_file_data = data
                self._imported_image_path = None
                self._imported_file_label.setText(f"📎 {fname} ({len(data)} wierszy)")
                self._imported_file_label.setStyleSheet("color: #00CC66; font-style: italic;")
                self._log(f"✓ Załadowano Word do AI: {fname} ({len(data)} wierszy)")
            elif ext in ('.jpg', '.jpeg', '.png', '.gif', '.webp', '.bmp'):
                self._imported_image_path = file_path
                self._imported_file_data = None
                size_kb = os.path.getsize(file_path) // 1024
                self._imported_file_label.setText(f"🖼 {fname} ({size_kb} KB)")
                self._imported_file_label.setStyleSheet("color: #00AAFF; font-style: italic;")
                self._log(f"✓ Załadowano zdjęcie do AI: {fname} ({size_kb} KB)")
            else:
                QMessageBox.warning(self, "Błąd", "Nieobsługiwany format pliku.")
        except Exception as e:
            self._log(f"❌ Błąd wczytania pliku: {e}")
            QMessageBox.critical(self, "Błąd", f"Nie udało się wczytać pliku:\n{e}")

    def _parse_excel(self, path):
        """Odczytaj wszystkie niepuste wiersze z aktywnego arkusza Excela"""
        wb = openpyxl.load_workbook(path, data_only=True)
        ws = wb.active
        rows = []
        for row in ws.iter_rows(values_only=True):
            row_vals = [str(cell) if cell is not None else "" for cell in row]
            if any(v.strip() for v in row_vals):
                rows.append(row_vals)
        return rows

    def _parse_word(self, path):
        """Odczytaj tabele (i ewentualnie akapity) z dokumentu Word"""
        doc = Document(path)
        rows = []
        for table in doc.tables:
            for row in table.rows:
                row_vals = [cell.text.strip() for cell in row.cells]
                if any(v for v in row_vals):
                    rows.append(row_vals)
        if not rows:
            for para in doc.paragraphs:
                if para.text.strip():
                    rows.append([para.text.strip()])
        return rows

    def eventFilter(self, obj, event):
        """Przechwytuj Backspace/Delete na QTableView — wyczyść zaznaczone komórki"""
        if obj is self.table_view and isinstance(event, QKeyEvent):
            if event.type() == event.Type.KeyPress and event.key() in (Qt.Key_Backspace, Qt.Key_Delete):
                selected = self.table_view.selectionModel().selectedIndexes()
                if selected:
                    old_state = self.model.get_state()
                    self.model.state_about_to_change.emit(old_state)
                    for idx in selected:
                        r, c = idx.row(), idx.column()
                        if 0 <= r < len(self.model.rows) and 0 <= c < len(self.model.rows[r]):
                            self.model.rows[r][c].value = ""
                    top_left = self.model.index(
                        min(i.row() for i in selected),
                        min(i.column() for i in selected))
                    bottom_right = self.model.index(
                        max(i.row() for i in selected),
                        max(i.column() for i in selected))
                    self.model.dataChanged.emit(top_left, bottom_right)
                    self.model.save_to_file()
                    self.model.state_changed.emit(self.model.get_state())
                    self._log(f"Wyczyszczono {len(selected)} komórek")
                    return True
        return super().eventFilter(obj, event)

    # --- AUTO-UPDATE HANDLERS (script-based) ---

    def check_for_updates_manual(self):
        """
        Pobierz nową wersję main.py z GitHub Raw (ręcznie, przez przycisk).
        Po pobraniu wystarczy zrestartować aplikację — bootstrap załaduje nowy skrypt.
        """
        if self._script_updater and self._script_updater.isRunning():
            return
        self._log("🔄 Sprawdzanie aktualizacji skryptu…")
        self.set_ai_status("Sprawdzanie…", "#FFA500")
        self._script_updater = ScriptUpdater()
        self._script_updater.updated.connect(self._on_script_updated)
        self._script_updater.up_to_date.connect(self._on_script_up_to_date)
        self._script_updater.failed.connect(self._on_script_update_failed)
        self._script_updater.start()

    def _on_script_updated(self, new_ver: str):
        self._log(f"✓ Zaktualizowano do v{new_ver} — zrestartuj aplikację, by wczytać zmiany.")
        self.set_ai_status("Aktualizacja gotowa — restart!", "#FFA500")
        self.statusBar().showMessage(
            f"ShiftFlow v{new_ver} gotowa — zamknij i otwórz ponownie.", 0)
        reply = QMessageBox.information(
            self,
            "Aktualizacja gotowa",
            f"Pobrano ShiftFlow v{new_ver}.\n"
            "Zamknij i uruchom aplikację ponownie, by zastosować zmiany.",
            QMessageBox.Ok,
        )

    def _on_script_up_to_date(self, ver: str):
        self._log(f"✓ Wersja v{ver} jest aktualna.")
        self.set_ai_status("AI: gotowe", "#00FF00")
        self.statusBar().showMessage(f"ShiftFlow v{ver} — brak aktualizacji.", 5000)

    def _on_script_update_failed(self, error: str):
        self._log(f"❌ Aktualizacja skryptu nie powiodła się: {error}")
        self.set_ai_status("AI: gotowe", "#00FF00")

    def _log(self, message):
        """Dodaj wiadomość do loga statusu"""
        current = self.status_log.toPlainText()
        self.status_log.setText(f"{current}\n{message}" if current else message)
        self.status_log.verticalScrollBar().setValue(
            self.status_log.verticalScrollBar().maximum()
        )

# --- PUNKT WEJŚCIA ---
if __name__ == "__main__":
    app = QApplication.instance() or QApplication(sys.argv)
    try:
        window = ScheduleApp()
        if _APP_CONFIG.get("start_maximized", True):
            window.showMaximized()
        else:
            window.show()
        sys.exit(app.exec())
    except Exception as exc:
        import traceback
        tb = traceback.format_exc()
        QMessageBox.critical(None, "Błąd krytyczny", f"{exc}\n\n{tb}")
        sys.exit(1)
